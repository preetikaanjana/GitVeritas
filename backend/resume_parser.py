import re
import os
from typing import List, Dict, Any, Tuple
from pypdf import PdfReader
from docx import Document

# A standard technology taxonomy to bootstrap direct search
TECH_TAXONOMY = {
    "languages": [
        "python", "javascript", "typescript", "java", "c++", "c#", "go", "golang", "rust",
        "ruby", "php", "swift", "kotlin", "scala", "r", "shell", "bash", "sql", "html", "css"
    ],
    "frameworks": [
        "react", "angular", "vue", "next.js", "nextjs", "nuxt", "svelte", "django", "flask",
        "fastapi", "express", "spring", "spring boot", "laravel", "rails", "ruby on rails",
        "asp.net", "dotnet", "dotnet core", "flutter", "react native", "electron"
    ],
    "libraries_tools": [
        "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn", "pandas", "numpy",
        "docker", "kubernetes", "k8s", "aws", "azure", "gcp", "google cloud", "terraform",
        "ansible", "jenkins", "git", "redis", "mongodb", "postgresql", "postgres", "mysql",
        "sqlite", "graphql", "rest api", "grpc", "webpack", "vite", "tailwind", "bootstrap"
    ]
}

class ResumeParser:
    def __init__(self):
        self.flat_taxonomy = []
        for category, skills in TECH_TAXONOMY.items():
            self.flat_taxonomy.extend(skills)
        self.flat_taxonomy = list(set(self.flat_taxonomy))

    def extract_text(self, file_path: str) -> str:
        """Extracts plain text from PDF or DOCX files."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self._extract_pdf_text(file_path)
        elif ext in ['.docx', '.doc']:
            return self._extract_docx_text(file_path)
        else:
            raise ValueError("Unsupported file format. Please upload PDF or DOCX.")

    def _extract_pdf_text(self, file_path: str) -> str:
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"Error reading PDF: {e}")
        return text

    def _extract_docx_text(self, file_path: str) -> str:
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"Error reading DOCX: {e}")
        return text

    def extract_skills(self, text: str) -> List[str]:
        """Extracts technical skills found directly matching the taxonomy."""
        found_skills = []
        text_lower = text.lower()
        
        # Word boundary search to prevent matches like 'c' in 'cat'
        for skill in self.flat_taxonomy:
            # Handle special characters in skills like c++, c#, next.js
            escaped_skill = re.escape(skill)
            pattern = rf"\b{escaped_skill}\b"
            if skill == "c++":
                pattern = r"c\+\+"
            elif skill == "c#":
                pattern = r"c\#"
            elif skill == "next.js":
                pattern = r"next\.js"
            elif skill == "dotnet":
                pattern = r"\.net"
            
            if re.search(pattern, text_lower):
                found_skills.append(skill)
        return found_skills

    def extract_quantifiable_claims(self, text: str) -> List[str]:
        """Extracts quantifiable bullet points related to software engineering projects and codebase metrics."""
        lines = text.split("\n")
        claims = []
        
        # Keywords to skip contact details and social media handles
        skip_keywords = [
            "email", "@", "phone", "mobile", "linkedin", "github", 
            "leetcode", "hackerrank", "address", "contact", "preetikaanjana"
        ]
        
        # Refined codebase engineering verbs (using word boundaries)
        engineering_verbs = [
            "built", "developed", "shipped", "designed", "implemented", "optimized", 
            "scaled", "integrated", "deployed", "migrated", "reduced", "increased", 
            "saved", "served", "wrote", "engineered", "created", "hosted"
        ]
        
        # Codebase and scale metrics
        metric_words = [
            "users", "percent", "%", "million", "thousand", "api", "queries", "requests", 
            "latency", "load", "speed", "database", "stars", "hours"
        ]
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            line_lower = line.lower()
            
            # 1. Skip contact details and profiles
            if any(skip in line_lower for skip in skip_keywords):
                continue
                
            # 2. Extract and check numbers (ignore 4-digit years like 2025, 2027)
            numbers = re.findall(r"\b\d+\b", line)
            has_metric_number = False
            for num in numbers:
                val = int(num)
                if 1990 <= val <= 2035:
                    continue  # Date year, not a metric
                has_metric_number = True
                break
                
            # Also match percentage signs or metrics suffixes (e.g. 50%, 200k, 10ms)
            if "%" in line or re.search(r"\b\d+(?:k|m|ms|fps|gb|mb|%)\b", line_lower):
                has_metric_number = True
                
            if not has_metric_number:
                continue
                
            # 3. Check for engineering verbs or metric words with strict boundaries
            has_verb = any(re.search(rf"\b{verb}\b", line_lower) for verb in engineering_verbs)
            has_metric = any(re.search(rf"\b{metric}\b", line_lower) for metric in metric_words)
            
            if has_verb or has_metric:
                # Clean line from leading bullet markers
                cleaned_line = re.sub(r"^[\s•\-*#\d\.]+", "", line).strip()
                if cleaned_line and len(cleaned_line) > 20:
                    claims.append(cleaned_line)
        
        return claims

    def extract_temporal_claims(self, text: str) -> Dict[str, Any]:
        """
        Parses chronological blocks to extract when skills were used.
        Returns a dict mapping skill to the years it was active on the resume.
        """
        # Split text into paragraphs/sections to isolate blocks of experience
        blocks = re.split(r"\n\s*\n|\n(?=[A-Z][a-z]+ \d{4})", text)
        skill_timelines = {}
        
        # Regex to find years (1990 - 2099)
        year_pattern = r"\b(19\d{2}|20\d{2})\b"
        # Range regex, e.g. "2020 - 2022" or "2021 to Present"
        range_pattern = r"\b(19\d{2}|20\d{2})\s*(?:-|to|–|—)\s*(Present|19\d{2}|20\d{2})\b"

        current_year = 2026 # Local system time reference is 2026

        for block in blocks:
            block = block.strip()
            if not block:
                continue
            
            # Look for date ranges in this block
            range_matches = re.findall(range_pattern, block, re.IGNORECASE)
            years = re.findall(year_pattern, block)
            
            start_year = None
            end_year = None
            
            if range_matches:
                # Use the first range found
                start_str, end_str = range_matches[0]
                start_year = int(start_str)
                if end_str.lower() == "present":
                    end_year = current_year
                else:
                    end_year = int(end_str)
            elif len(years) >= 2:
                # Fallback to min and max years in the block
                int_years = [int(y) for y in years]
                start_year = min(int_years)
                end_year = max(int_years)
            elif len(years) == 1:
                # Single year could mean a project or graduation year
                start_year = int(years[0])
                end_year = start_year
            
            if start_year and end_year:
                # Find skills mentioned in this block
                skills_in_block = self.extract_skills(block)
                for skill in skills_in_block:
                    if skill not in skill_timelines:
                        skill_timelines[skill] = []
                    
                    skill_timelines[skill].append({
                        "start_year": start_year,
                        "end_year": end_year,
                        "context": block[:150] + ("..." if len(block) > 150 else "")
                    })
        
        # Aggregate timelines per skill (take min start and max end)
        aggregated = {}
        for skill, periods in skill_timelines.items():
            min_start = min(p["start_year"] for p in periods)
            max_end = max(p["end_year"] for p in periods)
            contexts = [p["context"] for p in periods]
            aggregated[skill] = {
                "start_year": min_start,
                "end_year": max_end,
                "contexts": list(set(contexts))
            }
            
        return aggregated

    def parse(self, file_path: str) -> Dict[str, Any]:
        """Wrapper method to parse the entire resume."""
        text = self.extract_text(file_path)
        skills = self.extract_skills(text)
        quantifiable = self.extract_quantifiable_claims(text)
        timeline = self.extract_temporal_claims(text)
        
        return {
            "skills": skills,
            "quantifiable_claims": quantifiable,
            "timeline": timeline,
            "text_preview": text[:500] + "..." if len(text) > 500 else text
        }

if __name__ == "__main__":
    # Small test
    parser = ResumeParser()
    sample_text = """
    John Doe
    React Developer | Python Enthusiast
    
    Work Experience:
    Senior Web Developer at Acme Corp | Jan 2021 - Present
    - Built interactive user interfaces using React, TypeScript, and Tailwind CSS.
    - Led a team of 4 junior developers and optimized page load speed by 35%.
    - Formulated web APIs using Django and PostgreSQL.
    
    Data Engineer at Biotech Inc | 2018 - 2020
    - Managed big data pipelines using Python, pandas, and SQL databases.
    - Achieved 99.9% data reliability for 15,000 active users.
    """
    
    print("Direct Skills:", parser.extract_skills(sample_text))
    print("\nQuantifiable Claims:", parser.extract_quantifiable_claims(sample_text))
    print("\nTimeline:", parser.extract_temporal_claims(sample_text))
